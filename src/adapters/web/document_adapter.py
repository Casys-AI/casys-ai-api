import logging
from typing import List, Dict, Any, Union
from pathlib import Path
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document
from langchain.chains.summarize import load_summarize_chain

from src.domain.ports.document_adapter_protocol import DocumentAdapterProtocol
from src.infrastructure.config import config

logger = logging.getLogger("uvicorn.error")


class DocumentAdapter(DocumentAdapterProtocol):
    def __init__(self, openai_chat: Any):
        self.openai_chat = openai_chat
        self.text_splitter_config = config.get_text_splitter_config()
        self.text_splitter = self._create_text_splitter()
    
    def _create_text_splitter(self) -> RecursiveCharacterTextSplitter:
        return RecursiveCharacterTextSplitter(
            chunk_size=self.text_splitter_config["chunk_size"],
            chunk_overlap=self.text_splitter_config["chunk_overlap"],
            length_function=len,
            separators=self.text_splitter_config.get("separators", ["\n\n", "\n", " ", ""])
        )
    
    def load_document(self, file_path: str) -> str:
        logger.info(f"Starting document loading: {file_path}")
        file_path = Path(file_path)
        try:
            if file_path.suffix == ".pdf":
                content = self._load_pdf(file_path)
            elif file_path.suffix == ".docx":
                content = self._load_docx(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_path.suffix}")
            logger.info(f"Document loaded successfully: {file_path}")
            return content
        except Exception as e:
            logger.error(f"Error loading document {file_path}: {str(e)}")
            raise
    
    @staticmethod
    def _load_pdf(file_path: Path) -> str:
        with file_path.open("rb") as file:
            pdf_reader = PdfReader(file)
            return "\n".join(page.extract_text() for page in pdf_reader.pages)
    
    @staticmethod
    def _load_docx(file_path: Path) -> str:
        docx_doc = DocxDocument(file_path)
        return "\n".join(para.text for para in docx_doc.paragraphs)
    
    def split_text(self, text: str, chunk_size: int = None, chunk_overlap: int = None) -> List[Document]:
        logger.info("Starting text splitting")
        try:
            chunk_size = chunk_size or self.text_splitter_config["chunk_size"]
            chunk_overlap = chunk_overlap or self.text_splitter_config["chunk_overlap"]
            
            chunks = self.text_splitter.split_text(text)
            logger.info(f"Text split into {len(chunks)} chunks")
            return [Document(page_content=chunk) for chunk in chunks]
        except Exception as e:
            logger.error(f"Error during text splitting: {str(e)}")
            raise
    
    def process_document(self, doc: Union[Document, List[Document]]) -> Dict[str, str]:
        logger.info("Starting document processing")
        try:
            chain = load_summarize_chain(self.openai_chat, chain_type="stuff")
            result = chain.invoke([doc] if isinstance(doc, Document) else doc)
            logger.info("Document processed successfully")
            return {"summary": result['output_text'] if isinstance(result, dict) else result}
        except Exception as e:
            logger.error(f"Error during document processing: {str(e)}")
            raise
